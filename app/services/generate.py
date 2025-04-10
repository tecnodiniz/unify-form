import os

import re
import json
import uuid
import docx
import xlsxwriter
from datetime import datetime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from app.config import Config
from app.services import load_data, save_data


def generate_json(questions):
    client = OpenAI(api_key=Config.OPEN_IA_KEY)
    
    response = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {
        "role": "system",
        "content": "Você é um assistente que retorna apenas JSON válido. Sempre responda no formato `{ \"data\": [] }` sem explicações extras.\n\nVocê receberá um array contendo itens em string, onde cada item representa um formulário separado. Sua função é consolidar todos esses formulários em um único questionário inteligente unificado.\n\nCada item representa um questionário utilizado por seguradoras para cotação de um mesmo produto. Esses produtos estão agrupados em um dos seguintes **5 ramos principais de seguros**:\n- **Garantia**\n- **RC & RE (Responsabilidade Civil e Riscos de Engenharia)**\n- **Benefícios**\n- **Patrimonial**\n- **Especiais**\n\nCom base no conteúdo dos formulários, **identifique o ramo** e organize as perguntas de forma coerente, agrupando e padronizando os campos conforme as melhores práticas para aquele ramo.\n\nExemplos de seções típicas que podem aparecer:\n- Para **Garantia**: “Tomador”, “Objeto do Contrato”, “Valor Garantido”, “Prazo”, “Histórico Financeiro”\n- Para **RC & RE**: “Atividade”, “Faturamento”, “Exposição ao Risco”, “Coberturas”, “Reclamações Passadas”\n- Para **Benefícios**: “Funcionários”, “Plano de Saúde”, “Coberturas”, “Regras de Elegibilidade”\n- Para **Patrimonial**: “Local de Risco”, “Bens Segurados”, “Proteção e Segurança”, “Histórico de Sinistros”\n- Para **Especiais**: “Evento”, “Participantes”, “Infraestrutura”, “Coberturas Especiais”\n\nRetorne apenas JSON válido com a seguinte estrutura:\n`{ \"data\": [{ \"texto\": \"Questão\", \"tipo\": \"text\", \"obrigatorio\": \"false\", \"secao\": \"Geral\" }] }`\n\nA chave `\"texto\"` deve conter a pergunta padronizada em português (sem duplicações desnecessárias). \nA chave `\"tipo\"` deve ser uma destas: `\"text\"`, `\"number\"`, `\"date\"`, `\"select\"`, `\"boolean\"` ou `\"email\"` — escolha com base na natureza da pergunta.\nA chave `\"obrigatorio\"` deve conter **sempre uma string** com valor `\"true\"` ou `\"false\"`.\n\n⚠️ Se duas seguradoras perguntarem a mesma coisa com redações diferentes, normalize para uma só.\n⚠️ Se uma pergunta aparecer em apenas um formulário, ela ainda deve ser incluída.⚠️Local e Data: | Assinatura do Proponente: Você IGNORA. \n⚠️ Retorne apenas o JSON. Nenhuma explicação adicional."
        },
        {
            "role": "user",
            "content": json.dumps(questions)  # Enviando a lista como JSON válido
        }
    ],
    response_format={"type": "json_object"},  # Garantindo retorno JSON válido
    temperature=1,
    max_tokens=2048,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0
    )
    return response

def generate_unified_form(product, questions, insurers):
    """Gera um formulário unificado baseado nas perguntas extraídas"""
    form_id = str(uuid.uuid4())
    
    
    questions=generate_json(questions)
  
    json_load = json.loads(questions.choices[0].message.content)["data"]


    # Criar formulário unificado
    unified_form = {
        "id": form_id,
        "produto": product,
        "titulo": f"Questionário de cotação - {product}",
        "seguradoras": insurers,
        "perguntas": json_load,
        "data_criacao": datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    }

    # Salvar dados
    data = load_data()
    data["formularios"].append(unified_form)
    save_data(data)
    
    return unified_form

def generate_docx_form(form_data, output_path):
    """Gera um documento DOCX com o formulário unificado"""
    doc = docx.Document()

    def set_spacing(run, spacing_val):
        rPr = run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(spacing_val))
        rPr.append(spacing)

    # Header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    logo_path = os.path.abspath(os.path.join(Config.IMAGE_FOLDER, "logo.jpg"))

    if os.path.isfile(logo_path):
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(3))
        doc.add_paragraph()
    else:
        header_paragraph.text = "Facility & Bond Adm e Corretagem de Seguros Ltda."
        doc.add_paragraph()

    # Título
    doc.add_paragraph()
    doc.add_heading(form_data["titulo"], 0)

    doc.add_paragraph()

    sections = {}
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        if secao not in sections:
            sections[secao] = []
        sections[secao].append(pergunta)

    # Adicionar perguntas por seção
    for secao, perguntas in sections.items():
        if len(perguntas) > 0:
            doc.add_heading(secao, 1)
            for pergunta in perguntas:
                p = doc.add_paragraph()
                p.add_run(f"{pergunta['texto']}").bold = True
                if pergunta["obrigatorio"]:
                    p.add_run(" *").bold = True
                doc.add_paragraph("R:")
                doc.add_paragraph()

    # Rodapé
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 

    # Primeira linha com espaçamento expandido
    doc.add_paragraph()
    run1 = p.add_run("Facility & Bond Adm e Corretagem de Seguros Ltda.\n")
    set_spacing(run1, 40)  # 2pt = 40

    # Restante das informações do rodapé
    run2 = p.add_run(
        "Rua Pompeu Vairo, 123 - Vila Thaís\n"
        "CEP 12942 - 122 | Atibaia - SP\n"
        "Tel/Whatsapp: (11) 4418 - 8329\n"
        "e-mail: elementares@facilitybond.com.br"
    )
    run2.font.size = Pt(9)

    # Espaçamento antes da seção final
    doc.add_paragraph()

    # Criar uma tabela com 2 colunas para Local/Data e Assinatura
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3.5)

    # Primeira linha: rótulos
    cell1, cell2 = table.rows[0].cells
    cell1.paragraphs[0].add_run("Local e Data:").bold = True
    cell2.paragraphs[0].add_run("Assinatura do Proponente:").bold = True

    # Segunda linha: campos preenchíveis (linhas)
    cell1, cell2 = table.rows[1].cells
    cell1.paragraphs[0].add_run(" " * 40 + "___________________________")
    cell2.paragraphs[0].add_run(" " * 40 + "_____________________________")

    doc.save(output_path)
    return output_path

def generate_xlsx_form(form_data, output_path):
    """Gera uma planilha XLSX com o formulário unificado"""
    workbook = xlsxwriter.Workbook(output_path)
    worksheet = workbook.add_worksheet("Formulário")
    
    # Formatos
    title_format = workbook.add_format({
        'bold': True,
        'font_size': 16,
        'align': 'center',
        'valign': 'vcenter'
    })
    
    header_format = workbook.add_format({
        'bold': True,
        'bg_color': '#D9EAD3',
        'border': 1
    })
    
    section_format = workbook.add_format({
        'bold': True,
        'bg_color': '#F3F3F3',
        'border': 1
    })
    
    question_format = workbook.add_format({
        'text_wrap': True,
        'valign': 'top',
        'border': 1
    })
    
    answer_format = workbook.add_format({
        'bg_color': '#E9ECEF',
        'border': 1
    })
    
    # Título
    worksheet.merge_range('A1:C1', form_data["titulo"], title_format)
    worksheet.write('A2', f"Produto: {form_data['produto']}")
    worksheet.write('A3', f"Seguradoras: {', '.join(form_data['seguradoras'])}")
    worksheet.write('A4', f"Data de criação: {form_data['data_criacao']}")
    
    # Cabeçalhos
    worksheet.write('A6', "Seção", header_format)
    worksheet.write('B6', "Pergunta", header_format)
    worksheet.write('C6', "Resposta", header_format)
    
    # Ajustar larguras
    worksheet.set_column('A:A', 20)
    worksheet.set_column('B:B', 40)
    worksheet.set_column('C:C', 40)
    
    # Agrupar perguntas por seção
    sections = {
        "Dados do Proponente":[],
        "Informações do Evento":[],
        "Cobertura":[],
        "Histórico": [],
        "Informações Adicionais":[]
        }
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        # if secao not in sections:
        #     sections[secao] = []
        sections[secao].append(pergunta)
    
    # Adicionar perguntas por seção
    row = 6
    for secao, perguntas in sections.items():
        for i, pergunta in enumerate(perguntas):
            row += 1
            
            # Escrever seção apenas na primeira pergunta da seção
            if i == 0:
                worksheet.write(f'A{row}', secao, section_format)
            else:
                worksheet.write(f'A{row}', "", question_format)
            
            texto = pergunta["texto"]
            if pergunta["obrigatorio"]:
                texto += " *"
            
            worksheet.write(f'B{row}', texto, question_format)
            worksheet.write(f'C{row}', "", answer_format)
    
    # Rodapé
    row += 2
    worksheet.write(f'A{row}', f"Questionário de cotação - {form_data['produto']}")
    
    workbook.close()
    return output_path

def generate_html_form(form_data, output_path):
    """Gera um formulário HTML para preenchimento online"""
    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{form_data['titulo']}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {{
            font-family: 'Nunito', 'Segoe UI', Roboto, sans-serif;
            background-color: #f5f7fa;
            color: #333;
            line-height: 1.6;
        }}
        .form-container {{
            max-width: 800px;
            margin: 2rem auto;
            background-color: white;
            border-radius: 10px;
            box-shadow: 0 0.5rem 1rem rgba(0, 0, 0, 0.15);
            padding: 2rem;
        }}
        .form-header {{
            text-align: center;
            margin-bottom: 2rem;
            padding-bottom: 1rem;
            border-bottom: 1px solid #dee2e6;
        }}
        .form-title {{
            color: #0056b3;
            font-weight: 700;
        }}
        .form-section {{
            margin-bottom: 2rem;
            padding: 1rem;
            background-color: #f8f9fa;
            border-radius: 0.5rem;
            border-left: 4px solid #0056b3;
        }}
        .form-section-title {{
            font-weight: 700;
            color: #0056b3;
            margin-bottom: 1rem;
        }}
        .form-group {{
            margin-bottom: 1.5rem;
        }}
        .form-label {{
            font-weight: 600;
        }}
        .required {{
            color: #dc3545;
        }}
        .form-footer {{
            text-align: center;
            margin-top: 2rem;
            padding-top: 1rem;
            border-top: 1px solid #dee2e6;
            color: #6c757d;
            font-size: 0.9rem;
        }}
        .btn-primary {{
            background-color: #0056b3;
            border: none;
            padding: 0.6rem 1.5rem;
            font-weight: 600;
            border-radius: 50px;
        }}
        .btn-primary:hover {{
            background-color: #004494;
            transform: translateY(-2px);
            box-shadow: 0 5px 12px rgba(0, 0, 0, 0.15);
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="form-container">
            <div class="form-header">
                <h1 class="form-title">{form_data['titulo']}</h1>
                <p class="text-muted">Produto: {form_data['produto']} | Seguradoras: {', '.join(form_data['seguradoras'])}</p>
            </div>
            
            <form id="unified-form">
"""
    
    # Agrupar perguntas por seção
    sections = {}
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        if secao not in sections:
            sections[secao] = []
        sections[secao].append(pergunta)
    
    # Adicionar perguntas por seção
    for secao, perguntas in sections.items():
        html += f"""
                <div class="form-section">
                    <h3 class="form-section-title">{secao}</h3>
"""
        
        for pergunta in perguntas:
            pergunta_id = re.sub(r'[^\w]', '_', pergunta["texto"].lower())
            required = "required" if pergunta["obrigatorio"] else ""
            required_mark = '<span class="required">*</span>' if pergunta["obrigatorio"] else ""
            
            html += f"""
                    <div class="form-group">
                        <label for="{pergunta_id}" class="form-label">{pergunta['texto']} {required_mark}</label>
                        <input type="text" class="form-control" id="{pergunta_id}" name="{pergunta_id}" {required}>
                    </div>
"""
        
        html += """
                </div>
"""
    
    # Finalizar HTML
    html += """
                <div class="d-grid gap-2">
                    <button type="submit" class="btn btn-primary">Enviar Formulário</button>
                </div>
            </form>
            
            <div class="form-footer">
                <p>Este formulário unificado foi gerado automaticamente e contém perguntas de múltiplas seguradoras.</p>
            </div>
        </div>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        document.getElementById('unified-form').addEventListener('submit', function(e) {
            e.preventDefault();
            alert('Formulário enviado com sucesso!');
        });
    </script>
</body>
</html>
"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    
    return output_path
