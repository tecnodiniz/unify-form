from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
import os
import json
import uuid
from datetime import datetime
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import docx
from docx.shared import Inches , Pt
import re
import xlsxwriter
import pdb # Debuger
from werkzeug.utils import secure_filename
from openai import OpenAI
from dotenv import load_dotenv
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

api_key=os.getenv("OPENAI_API_KEY")
app = Flask(__name__)
app.secret_key = os.urandom(24)
app.jinja_env.globals.update(now=datetime.now)

# Configurações
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
FORMS_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'forms')
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
DATA_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data.json')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['FORMS_FOLDER'] = FORMS_FOLDER

# Garantir que os diretórios existam
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(FORMS_FOLDER, exist_ok=True)

# Funções auxiliares
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_data():
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {"produtos": [], "formularios": []}

def save_data(data):
    with open(DATA_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def extract_questions_from_pdf(file_path):
    """Extrai perguntas de um arquivo PDF"""
    questions = []
    try:
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            for page_num in range(len(pdf.pages)):
                page = pdf.pages[page_num]
                text = page.extract_text()
                questions.append(text)
          
                # # Procurar por padrões comuns de perguntas
                # patterns = [
                #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\?',  # Texto seguido por ponto de interrogação
                #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+):\s*',  # Texto seguido por dois pontos
                #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\s*_+',  # Texto seguido por sublinhados
                #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\[\s*\]'  # Texto seguido por colchetes vazios
                # ]
                
                # for pattern in patterns:
                #     matches = re.findall(pattern, text)
                #     for match in matches:
                #         question = match.strip()
                #         if question and len(question) > 3 and question not in [q["texto"] for q in questions]:
                #             questions.append({
                #                 "texto": question,
                #                 "tipo": "text",
                #                 "obrigatorio": False,
                #                 "secao": "Geral"
                #             })
    except Exception as e:
        print(f"Erro ao extrair perguntas do PDF: {e}")
    
    return questions

def extract_questions_from_docx(file_path):
    """Extrai perguntas de um arquivo DOCX"""
    questions = []
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        questions.append(text)
     
        # # Procurar por padrões comuns de perguntas
        # patterns = [
        #     r'/^.(?:\d)?+(?:\.\s*)?[A-Za-zÀ-ÖØ-öø-ÿ\s\d\(\) \/,]+[:?]\s*$/gm'
        #     r'^\d+(?:\.\s*)?[A-Za-zÀ-ÖØ-öø-ÿ\s\d\(\),]+:\s*$',
        #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\?',  # Texto seguido por ponto de interrogação
        #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+):\s*'  # Texto seguido por dois pontos
        #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\s*_+',  # Texto seguido por sublinhados
        #     r'([A-Za-zÀ-ÖØ-öø-ÿ\s\d]+)\[\s*\]',  # Texto seguido por colchetes vazios
        #     r'\d+\.\s*[A-Za-zÀ-ÖØ-öø-ÿ\s\d]+:'
        # ]
        
        # for pattern in patterns:
        #     matches = re.findall(pattern, text)
        #     for match in matches:
        #         question = match.strip()
        #         if question and len(question) > 3 and question not in [q["texto"] for q in questions]:
        #             questions.append({
        #                 "texto": question,
        #                 "tipo": "text",
        #                 "obrigatorio": False,
        #                 "secao": "Geral"
        #             })
    except Exception as e:
        print(f"Erro ao extrair perguntas do DOCX: {e}")
    
   
    return questions

def extract_questions(file_path):
    """Extrai perguntas de um arquivo baseado em sua extensão"""
    ext = file_path.rsplit('.', 1)[1].lower()
    
    if ext == 'pdf':
        return extract_questions_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        return extract_questions_from_docx(file_path)
    else:
        return []

def organize_questions_by_section(questions, product):
    """Organiza perguntas em seções lógicas baseadas em palavras-chave"""
    sections = {}

    match product:
        case "rc_eventos":
            sections = {
                "Dados do Proponente": ["nome", "cnpj", "cpf", "endereço", "telefone", "email", "contato", "empresa", "proponente"],
                "Informações do Evento": ["evento", "data", "local", "participantes", "público", "duração", "horário"],
                "Cobertura": ["cobertura", "limite", "valor", "indenização", "franquia", "seguro", "apólice"],
                "Histórico": ["sinistro", "ocorrência", "reclamação", "histórico", "anterior"],
                "Informações Adicionais": []
            }
    
    organized_questions = []
    

    for question in questions:
        assigned = False
        question_lower = question["texto"].lower()
        
        for section, keywords in sections.items():
            for keyword in keywords:
                if keyword in question_lower:
                    question["secao"] = section
                    organized_questions.append(question)
                    assigned = True
                    break
            if assigned:
                break
        
        if not assigned:
            question["secao"] = "Informações Adicionais"
            organized_questions.append(question)
    
    return organized_questions

# def remove_duplicate_questions(questions):
#     """Remove perguntas duplicadas ou muito similares"""
#     unique_questions = []
#     seen_texts = set()
    
#     for question in questions:
#         # Normalizar texto para comparação
#         normalized_text = re.sub(r'[^\w\s]', '', question["texto"].lower())
#         normalized_text = re.sub(r'\s+', ' ', normalized_text).strip()
        
#         # Verificar se já existe uma pergunta similar
#         is_duplicate = False
#         for seen_text in seen_texts:
#             # Calcular similaridade (implementação simples)
#             if normalized_text in seen_text or seen_text in normalized_text:
#                 is_duplicate = True
#                 break
            
#             # Outra opção: calcular distância de Levenshtein
#             # Se a similaridade for alta, considerar como duplicata
        
#         if not is_duplicate:
#             seen_texts.add(normalized_text)
#             unique_questions.append(question)
    
#     return unique_questions
def ai_filter(questions):
    client = OpenAI(api_key=api_key)
    
    response = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {
        "role": "system",
        "content": "Você é um assistente que retorna apenas JSON válido. Sempre responda no formato `{ \"data\": [] }` sem explicações extras.\n\nVocê receberá um array contendo itens em string, onde cada item representa um formulário separado. Sua função é consolidar todos esses formulários em um único questionário inteligente unificado.\n\nCada item representa um questionário utilizado por seguradoras para cotação de um mesmo produto. Esses produtos estão agrupados em um dos seguintes **5 ramos principais de seguros**:\n- **Garantia**\n- **RC & RE (Responsabilidade Civil e Riscos de Engenharia)**\n- **Benefícios**\n- **Patrimonial**\n- **Especiais**\n\nCom base no conteúdo dos formulários, **identifique o ramo** e organize as perguntas de forma coerente, agrupando e padronizando os campos conforme as melhores práticas para aquele ramo.\n\nExemplos de seções típicas que podem aparecer:\n- Para **Garantia**: “Tomador”, “Objeto do Contrato”, “Valor Garantido”, “Prazo”, “Histórico Financeiro”\n- Para **RC & RE**: “Atividade”, “Faturamento”, “Exposição ao Risco”, “Coberturas”, “Reclamações Passadas”\n- Para **Benefícios**: “Funcionários”, “Plano de Saúde”, “Coberturas”, “Regras de Elegibilidade”\n- Para **Patrimonial**: “Local de Risco”, “Bens Segurados”, “Proteção e Segurança”, “Histórico de Sinistros”\n- Para **Especiais**: “Evento”, “Participantes”, “Infraestrutura”, “Coberturas Especiais”\n\nRetorne apenas JSON válido com a seguinte estrutura:\n`{ \"data\": [{ \"texto\": \"Questão\", \"tipo\": \"text\", \"obrigatorio\": \"false\", \"secao\": \"Geral\" }] }`\n\nA chave `\"texto\"` deve conter a pergunta padronizada em português (sem duplicações desnecessárias). \nA chave `\"tipo\"` deve ser uma destas: `\"text\"`, `\"number\"`, `\"date\"`, `\"select\"`, `\"boolean\"` ou `\"email\"` — escolha com base na natureza da pergunta.\nA chave `\"obrigatorio\"` deve conter **sempre uma string** com valor `\"true\"` ou `\"false\"`.\n\n⚠️ Se duas seguradoras perguntarem a mesma coisa com redações diferentes, normalize para uma só.\n⚠️ Se uma pergunta aparecer em apenas um formulário, ela ainda deve ser incluída.⚠️Local e Data: | Assinatura do Proponente: Você IGNORA. \n⚠️ Retorne apenas o JSON. Nenhuma explicação adicional."
        },
        {
            "role": "user",
            "content": json.dumps(questions)  # Enviando a lista como JSON válido
        }
    ],
    response_format={"type": "json_object"},  # Garantindo retorno JSON válido
    temperature=1,
    max_tokens=2048,
    top_p=1,
    frequency_penalty=0,
    presence_penalty=0
    )
    return response
def generate_unified_form(product, questions, insurers):
    """Gera um formulário unificado baseado nas perguntas extraídas"""
    form_id = str(uuid.uuid4())
    
    # Organizar perguntas por seção

    # Descomentar para produção
    questions=ai_filter(questions)
  
    json_load = json.loads(questions.choices[0].message.content)["data"]
    # organized_questions = organize_questions_by_section(json_load, product)
    #------------------------------------------------------------------------
  

    # Criar formulário unificado
    unified_form = {
        "id": form_id,
        "produto": product,
        "titulo": f"Questionário de cotação - {product}",
        "seguradoras": insurers,
        "perguntas": json_load,
        "data_criacao": datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    }

    # Salvar dados
    data = load_data()
    data["formularios"].append(unified_form)
    save_data(data)
    
    return unified_form

def generate_docx_form(form_data, output_path):
    """Gera um documento DOCX com o formulário unificado"""
    doc = docx.Document()

    def set_spacing(run, spacing_val):
        rPr = run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(spacing_val))
        rPr.append(spacing)

    # Header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    logo_path = os.path.abspath(os.path.join("img", "logo.jpg"))

    if os.path.isfile(logo_path):
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(3))
    else:
        header_paragraph.text = "Facility & Bond Adm e Corretagem de Seguros Ltda."

    # Título
    doc.add_paragraph()
    doc.add_heading(form_data["titulo"], 0)

    doc.add_paragraph()

    # Agrupar perguntas por seção
    # sections = {
    #     "Dados do Proponente": [],
    #     "Informações do Evento": [],
    #     "Cobertura": [],
    #     "Histórico": [],
    #     "Informações Adicionais": []
    # }
    sections = {}
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        if secao not in sections:
            sections[secao] = []
        sections[secao].append(pergunta)

    # Adicionar perguntas por seção
    for secao, perguntas in sections.items():
        if len(perguntas) > 0:
            doc.add_heading(secao, 1)
            for pergunta in perguntas:
                p = doc.add_paragraph()
                p.add_run(f"{pergunta['texto']}").bold = True
                if pergunta["obrigatorio"]:
                    p.add_run(" *").bold = True
                doc.add_paragraph("R:")
                doc.add_paragraph()

    # Rodapé
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
 

    # Primeira linha com espaçamento expandido
    run1 = p.add_run("Facility & Bond Adm e Corretagem de Seguros Ltda.\n")
    set_spacing(run1, 40)  # 2pt = 40

    # Restante das informações do rodapé
    run2 = p.add_run(
        "Rua pompeu Vairo, 123 - Vila Thaís\n"
        "CEP 12942 - 122 | Atibaia - SP\n"
        "Tel/Whatsapp: (11) 4418 - 8329\n"
        "e-mail: elementares@facilitybond.com.br"
    )
    run2.font.size = Pt(9)

    # Espaçamento antes da seção final
    doc.add_paragraph()

    # Criar uma tabela com 2 colunas para Local/Data e Assinatura
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3.5)

    # Primeira linha: rótulos
    cell1, cell2 = table.rows[0].cells
    cell1.paragraphs[0].add_run("Local e Data:").bold = True
    cell2.paragraphs[0].add_run("Assinatura do Proponente:").bold = True

    # Segunda linha: campos preenchíveis (linhas)
    cell1, cell2 = table.rows[1].cells
    cell1.paragraphs[0].add_run(" " * 40 + "___________________________")
    cell2.paragraphs[0].add_run(" " * 40 + "_____________________________")

    doc.save(output_path)
    return output_path

def generate_xlsx_form(form_data, output_path):
    """Gera uma planilha XLSX com o formulário unificado"""
    workbook = xlsxwriter.Workbook(output_path)
    worksheet = workbook.add_worksheet("Formulário")
    
    # Formatos
    title_format = workbook.add_format({
        'bold': True,
        'font_size': 16,
        'align': 'center',
        'valign': 'vcenter'
    })
    
    header_format = workbook.add_format({
        'bold': True,
        'bg_color': '#D9EAD3',
        'border': 1
    })
    
    section_format = workbook.add_format({
        'bold': True,
        'bg_color': '#F3F3F3',
        'border': 1
    })
    
    question_format = workbook.add_format({
        'text_wrap': True,
        'valign': 'top',
        'border': 1
    })
    
    answer_format = workbook.add_format({
        'bg_color': '#E9ECEF',
        'border': 1
    })
    
    # Título
    worksheet.merge_range('A1:C1', form_data["titulo"], title_format)
    worksheet.write('A2', f"Produto: {form_data['produto']}")
    worksheet.write('A3', f"Seguradoras: {', '.join(form_data['seguradoras'])}")
    worksheet.write('A4', f"Data de criação: {form_data['data_criacao']}")
    
    # Cabeçalhos
    worksheet.write('A6', "Seção", header_format)
    worksheet.write('B6', "Pergunta", header_format)
    worksheet.write('C6', "Resposta", header_format)
    
    # Ajustar larguras
    worksheet.set_column('A:A', 20)
    worksheet.set_column('B:B', 40)
    worksheet.set_column('C:C', 40)
    
    # Agrupar perguntas por seção
    sections = {
        "Dados do Proponente":[],
        "Informações do Evento":[],
        "Cobertura":[],
        "Histórico": [],
        "Informações Adicionais":[]
        }
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        # if secao not in sections:
        #     sections[secao] = []
        sections[secao].append(pergunta)
    
    # Adicionar perguntas por seção
    row = 6
    for secao, perguntas in sections.items():
        for i, pergunta in enumerate(perguntas):
            row += 1
            
            # Escrever seção apenas na primeira pergunta da seção
            if i == 0:
                worksheet.write(f'A{row}', secao, section_format)
            else:
                worksheet.write(f'A{row}', "", question_format)
            
            texto = pergunta["texto"]
            if pergunta["obrigatorio"]:
                texto += " *"
            
            worksheet.write(f'B{row}', texto, question_format)
            worksheet.write(f'C{row}', "", answer_format)
    
    # Rodapé
    row += 2
    worksheet.write(f'A{row}', f"Questionário de cotação - {form_data['produto']}")
    
    workbook.close()
    return output_path

def generate_html_form(form_data, output_path):
    """Gera um formulário HTML para preenchimento online"""
    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{form_data['titulo']}</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        body {{
            font-family: 'Nunito', 'Segoe UI', Roboto, sans-serif;
            background-color: #f5f7fa;
            color: #333;
            line-height: 1.6;
        }}
        .form-container {{
            max-width: 800px;
            margin: 2rem auto;
            background-color: white;
            border-radius: 10px;
            box-shadow: 0 0.5rem 1rem rgba(0, 0, 0, 0.15);
            padding: 2rem;
        }}
        .form-header {{
            text-align: center;
            margin-bottom: 2rem;
            padding-bottom: 1rem;
            border-bottom: 1px solid #dee2e6;
        }}
        .form-title {{
            color: #0056b3;
            font-weight: 700;
        }}
        .form-section {{
            margin-bottom: 2rem;
            padding: 1rem;
            background-color: #f8f9fa;
            border-radius: 0.5rem;
            border-left: 4px solid #0056b3;
        }}
        .form-section-title {{
            font-weight: 700;
            color: #0056b3;
            margin-bottom: 1rem;
        }}
        .form-group {{
            margin-bottom: 1.5rem;
        }}
        .form-label {{
            font-weight: 600;
        }}
        .required {{
            color: #dc3545;
        }}
        .form-footer {{
            text-align: center;
            margin-top: 2rem;
            padding-top: 1rem;
            border-top: 1px solid #dee2e6;
            color: #6c757d;
            font-size: 0.9rem;
        }}
        .btn-primary {{
            background-color: #0056b3;
            border: none;
            padding: 0.6rem 1.5rem;
            font-weight: 600;
            border-radius: 50px;
        }}
        .btn-primary:hover {{
            background-color: #004494;
            transform: translateY(-2px);
            box-shadow: 0 5px 12px rgba(0, 0, 0, 0.15);
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="form-container">
            <div class="form-header">
                <h1 class="form-title">{form_data['titulo']}</h1>
                <p class="text-muted">Produto: {form_data['produto']} | Seguradoras: {', '.join(form_data['seguradoras'])}</p>
            </div>
            
            <form id="unified-form">
"""
    
    # Agrupar perguntas por seção
    sections = {}
    for pergunta in form_data["perguntas"]:
        secao = pergunta["secao"]
        if secao not in sections:
            sections[secao] = []
        sections[secao].append(pergunta)
    
    # Adicionar perguntas por seção
    for secao, perguntas in sections.items():
        html += f"""
                <div class="form-section">
                    <h3 class="form-section-title">{secao}</h3>
"""
        
        for pergunta in perguntas:
            pergunta_id = re.sub(r'[^\w]', '_', pergunta["texto"].lower())
            required = "required" if pergunta["obrigatorio"] else ""
            required_mark = '<span class="required">*</span>' if pergunta["obrigatorio"] else ""
            
            html += f"""
                    <div class="form-group">
                        <label for="{pergunta_id}" class="form-label">{pergunta['texto']} {required_mark}</label>
                        <input type="text" class="form-control" id="{pergunta_id}" name="{pergunta_id}" {required}>
                    </div>
"""
        
        html += """
                </div>
"""
    
    # Finalizar HTML
    html += """
                <div class="d-grid gap-2">
                    <button type="submit" class="btn btn-primary">Enviar Formulário</button>
                </div>
            </form>
            
            <div class="form-footer">
                <p>Este formulário unificado foi gerado automaticamente e contém perguntas de múltiplas seguradoras.</p>
            </div>
        </div>
    </div>
    
    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        document.getElementById('unified-form').addEventListener('submit', function(e) {
            e.preventDefault();
            alert('Formulário enviado com sucesso!');
        });
    </script>
</body>
</html>
"""
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    
    return output_path

# Rotas
@app.route('/')
def index():
    data = load_data()
    return render_template('index.html', data=data)

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'POST':
        if 'files[]' not in request.files:
            flash('Nenhum arquivo selecionado')
            return redirect(request.url)
        
        files = request.files.getlist('files[]')
        product = request.form.get('produto', '')
        
       
  
        if not product:
            flash('Nome do produto é obrigatório')
            return redirect(request.url)
        
        if not files or files[0].filename == '':
            flash('Nenhum arquivo selecionado')
            return redirect(request.url)
        
        # Verificar se todos os arquivos são permitidos
        for file in files:
            if not allowed_file(file.filename):
                flash(f'Tipo de arquivo não permitido: {file.filename}')
                return redirect(request.url)
        
        # Processar arquivos
        all_questions = []
        insurers = []
        
   
        for file in files:
            filename = secure_filename(file.filename)
            insurer = request.form.get(f'seguradora_{files.index(file)}', 'Não especificada')
            insurers.append(insurer)
            
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uuid.uuid4()}_{filename}")
            file.save(file_path)
            
            # Extrair perguntas
          
            questions = extract_questions(file_path)
       
            all_questions.extend(questions)
        
        # Gerar formulário unificado
        
        unified_form = generate_unified_form(product, all_questions, insurers)
    
     
        
        # Gerar arquivos para download
        form_id = unified_form["id"]
        docx_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.docx")
        xlsx_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.xlsx")
        html_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.html")
        
        generate_docx_form(unified_form, docx_path)
        # generate_xlsx_form(unified_form, xlsx_path)
        generate_html_form(unified_form, html_path)
        
        # Adicionar produto à lista se não existir
        data = load_data()
        if product not in [p for p in data["produtos"]]:
            data["produtos"].append(product)
            save_data(data)
        
        return redirect(url_for('view_form', form_id=form_id))
    
    return render_template('upload.html')

@app.route('/form/<form_id>')
def view_form(form_id):
    data = load_data()
    form = next((f for f in data["formularios"] if f["id"] == form_id), None)
    
    if not form:
        flash('Formulário não encontrado')
        return redirect(url_for('index'))
    
    return render_template('view_form.html', form=form)

@app.route('/download/<form_id>/<format>')
def download_form(form_id, format):
    data = load_data()
    form = next((f for f in data["formularios"] if f["id"] == form_id), None)
    
    if not form:
        flash('Formulário não encontrado')
        return redirect(url_for('index'))
    
    if format == 'docx':
        file_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.docx")
        return send_file(file_path, as_attachment=True, download_name=f"formulario_{form['produto']}.docx")
    
    elif format == 'xlsx':
        file_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.xlsx")
        return send_file(file_path, as_attachment=True, download_name=f"formulario_{form['produto']}.xlsx")
    
    elif format == 'html':
        file_path = os.path.join(app.config['FORMS_FOLDER'], f"{form_id}.html")
        return send_file(file_path, as_attachment=True, download_name=f"formulario_{form['produto']}.html")
    
    else:
        flash('Formato não suportado')
        return redirect(url_for('view_form', form_id=form_id))

@app.route('/online/<form_id>')
def online_form(form_id):
    data = load_data()
    form = next((f for f in data["formularios"] if f["id"] == form_id), None)
    
    if not form:
        flash('Formulário não encontrado')
        return redirect(url_for('index'))
    
    return render_template('online_form.html', form=form)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
